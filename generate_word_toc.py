import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_word_toc():
    # Create a new Word document
    doc = Document()
    
    # Set margins (narrower than default)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add the header - Table of Contents
    header = doc.add_heading('TABLE OF CONTENTS', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing after the header
    doc.add_paragraph()
    
    # Add the main sections
    sections = [
        ("ABSTRACT", "iv"),
        ("TABLE OF CONTENTS", "v"),
        ("LIST OF FIGURES", "viii"),
        ("LIST OF TABLES", "x")
    ]
    
    # Add the main sections
    for title, page in sections:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        
        # Add tab and page number aligned to the right
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        p.add_run("\t" + page)
    
    # Add spacing before chapters
    doc.add_paragraph()
    
    # Add the chapter headings
    p = doc.add_paragraph()
    p.add_run("CHAPTER NO.").bold = True
    p.add_run("\t").bold = True
    p.add_run("TITLE").bold = True
    
    # Add tab and page number aligned to the right
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run("\t" + "PAGE NO.").bold = True
    
    # Chapter 1 content
    add_chapter(doc, "1", "INTRODUCTION", "1")
    add_section(doc, "1.1", "Introduction to Project", "1")
    add_section(doc, "1.2", "Motivation", "2")
    add_section(doc, "1.3", "Sustainable Development Goal of the Project", "3")
    add_section(doc, "1.4", "Product Vision Statement", "4")
    add_section(doc, "1.5", "Product Goal", "5")
    add_section(doc, "1.6", "Product Backlog (Key User Stories with Desired Outcomes)", "6")
    add_section(doc, "1.7", "Product Release Plan", "8")
    
    # Add spacing
    doc.add_paragraph()
    
    # Chapter 2 content
    add_chapter(doc, "2", "SPRINT PLANNING AND EXECUTION", "9")
    add_section(doc, "2.1", "Sprint 1", "9")
    add_subsection(doc, "2.1.1", "Sprint Goal with User Stories of Sprint 1", "9")
    add_subsection(doc, "2.1.2", "Functional Document", "13")
    add_subsection(doc, "2.1.3", "Architecture Document", "16")
    add_subsection(doc, "2.1.4", "UI Design", "18")
    add_subsection(doc, "2.1.5", "Functional Test Cases", "19")
    add_subsection(doc, "2.1.6", "Daily Call Progress", "19")
    add_subsection(doc, "2.1.7", "Committed vs Completed User Stories", "20")
    add_subsection(doc, "2.1.8", "Sprint Retrospective", "20")
    
    # Add spacing
    doc.add_paragraph()
    
    # Section 2.2
    add_section(doc, "2.2", "Sprint 2", "21")
    add_subsection(doc, "2.2.1", "Sprint Goal with User Stories of Sprint 2", "21")
    add_subsection(doc, "2.2.2", "Functional Document", "25")
    add_subsection(doc, "2.2.3", "Architecture Document", "28")
    add_subsection(doc, "2.2.4", "UI Design", "30")
    add_subsection(doc, "2.2.5", "Functional Test Cases", "31")
    add_subsection(doc, "2.2.6", "Daily Call Progress", "31")
    add_subsection(doc, "2.2.7", "Committed vs Completed User Stories", "32")
    add_subsection(doc, "2.2.8", "Sprint Retrospective", "32")
    
    # Add spacing
    doc.add_paragraph()
    
    # Section 2.3
    add_section(doc, "2.3", "Sprint 3", "33")
    add_subsection(doc, "2.3.1", "Sprint Goal with User Stories of Sprint 3", "33")
    add_subsection(doc, "2.3.2", "Functional Document", "37")
    add_subsection(doc, "2.3.3", "Architecture Document", "40")
    add_subsection(doc, "2.3.4", "UI Design", "42")
    add_subsection(doc, "2.3.5", "Functional Test Cases", "44")
    add_subsection(doc, "2.3.6", "Daily Call Progress", "44")
    add_subsection(doc, "2.3.7", "Committed vs Completed User Stories", "45")
    add_subsection(doc, "2.3.8", "Sprint Retrospective", "45")
    
    # Add spacing
    doc.add_paragraph()
    
    # Chapter 3 content
    add_chapter(doc, "3", "RESULTS AND DISCUSSIONS", "46")
    add_section(doc, "3.1", "Project Outcomes", "46")
    add_section(doc, "3.2", "Committed vs Completed User Stories", "47")
    
    # Add spacing
    doc.add_paragraph()
    
    # Chapter 4 content
    add_chapter(doc, "4", "CONCLUSIONS & FUTURE ENHANCEMENT", "48")
    
    # Add spacing
    doc.add_paragraph()
    
    # References and Appendix
    p = doc.add_paragraph()
    p.add_run("REFERENCES").bold = True
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run("\t49")
    
    # Add spacing
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("APPENDIX").bold = True
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run("\t50")
    
    add_section(doc, "A.", "Sample Coding", "50")
    
    # Add spacing
    doc.add_paragraph()
    
    # LIST OF FIGURES
    doc.add_heading('LIST OF FIGURES', level=1)
    p = doc.add_paragraph()
    p.add_run("FIG NO.").bold = True
    p.add_run("\t").bold = True
    p.add_run("TITLE").bold = True
    
    # Add tab and page number aligned to the right
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run("\t" + "PAGE NO.").bold = True
    
    # Figures content
    figures = [
        ("1.1", "MS Planner Board of NutriFit Analytics and Dashboard System", "7"),
        ("1.2", "Release Plan of NutriFit Analytics and Dashboard System", "8"),
        ("2.1", "User Story for Key Metrics Dashboard", "10"),
        ("2.2", "User Story for Nutrition Data Filtering Panel", "11"),
        ("2.3", "User Story for Regional Distribution Treemap", "12"),
        ("2.4", "System Architecture Diagram for Sprint 1", "16"),
        ("2.5", "UI Design for Nutrition Overview", "18"),
        ("2.6", "UI Design for Nutrition Overview with Filter Options", "18"),
        ("2.7", "UI Design for Regional Analysis", "18"),
        ("2.8", "Standup Meetings for Sprint 1", "19"),
        ("2.9", "Bar Graph for Committed vs Completed User Stories for Sprint 1", "20"),
        ("2.10", "User Story for Registration Form for New Users", "22"),
        ("2.11", "User Story for Password Reset Workflow", "23"),
        ("2.12", "User Story for Admin Dashboard with Advanced Analytics", "24"),
        ("2.13", "System Architecture Diagram for Sprint 2", "28"),
        ("2.14", "UI Design for Login Page", "30"),
        ("2.15", "UI Design for Signup Page", "30"),
        ("2.16", "UI Design for Advanced Analytics Charts", "30"),
        ("2.17", "Standup Meetings for Sprint 2", "31"),
        ("2.18", "Bar Graph for Committed vs Completed User Stories for Sprint 2", "32"),
        ("2.19", "User Story for Nutrition Record Addition Functionality", "34"),
        ("2.20", "User Story for Nutrition Record Update Functionality", "35"),
        ("2.21", "User Story for Nutrition Record Deletion Functionality", "36"),
        ("2.22", "System Architecture Diagram for Sprint 3", "41"),
        ("2.23", "UI Design for Adding a Record", "42"),
        ("2.24", "UI Design for Data Management Page", "43"),
        ("2.25", "UI Design for Deleting a Record", "43"),
        ("2.26", "Standup Meetings for Sprint 3", "44"),
        ("2.27", "Bar Graph for Committed vs Completed User Stories for Sprint 3", "45"),
        ("3.1", "Committed vs Completed User Stories for All Sprints", "47")
    ]
    
    for num, title, page in figures:
        add_item(doc, num, title, page)
    
    # Add spacing
    doc.add_paragraph()
    
    # LIST OF TABLES
    doc.add_heading('LIST OF TABLES', level=1)
    p = doc.add_paragraph()
    p.add_run("TABLE NO.").bold = True
    p.add_run("\t").bold = True
    p.add_run("TITLE").bold = True
    
    # Add tab and page number aligned to the right
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run("\t" + "PAGE NO.").bold = True
    
    # Tables content
    tables = [
        ("1.1", "Product Backlog of NutriFit Analytics and Dashboard System", "6"),
        ("2.1", "Detailed User Stories of Sprint 1", "9"),
        ("2.2", "Access Level Authorization Matrix for Sprint 1", "15"),
        ("2.3", "Detailed Functional Test Case for Sprint 1", "19"),
        ("2.4", "Sprint Retrospective for Sprint 1", "20"),
        ("2.5", "Detailed User Stories of Sprint 2", "21"),
        ("2.6", "Access Level Authorization Matrix for Sprint 2", "27"),
        ("2.7", "Detailed Functional Test Case for Sprint 2", "31"),
        ("2.8", "Sprint Retrospective for Sprint 2", "32"),
        ("2.9", "Detailed User Stories of Sprint 3", "33"),
        ("2.10", "Access Level Authorization Matrix for Sprint 3", "39"),
        ("2.11", "Detailed Functional Test Case for Sprint 3", "44"),
        ("2.12", "Sprint Retrospective for Sprint 3", "45")
    ]
    
    for num, title, page in tables:
        add_item(doc, num, title, page)
    
    # Save the document
    doc.save('NutriFit_TOC_Updated.docx')
    print("Word document created successfully as 'NutriFit_TOC_Updated.docx'")

def add_chapter(doc, number, title, page):
    p = doc.add_paragraph()
    p.add_run(f"{number}").bold = True
    
    # Add tab for title with bold
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run(f"\t{title}").bold = True
    
    # Add tab for page number
    p.add_run(f"\t{page}")

def add_section(doc, number, title, page):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(f"{number}")
    
    # Add tab for title
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run(f"\t{title}")
    
    # Add tab for page number
    p.add_run(f"\t{page}")

def add_subsection(doc, number, title, page):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.6)
    p.add_run(f"{number}")
    
    # Add tab for title
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run(f"\t{title}")
    
    # Add tab for page number
    p.add_run(f"\t{page}")

def add_item(doc, number, title, page):
    p = doc.add_paragraph()
    p.add_run(f"{number}")
    
    # Add tab for title
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run(f"\t{title}")
    
    # Add tab for page number
    p.add_run(f"\t{page}")

if __name__ == "__main__":
    generate_word_toc() 