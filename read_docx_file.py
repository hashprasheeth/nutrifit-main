import docx
import os

def extract_toc_from_docx(docx_path):
    print(f"Reading document: {docx_path}")
    try:
        doc = docx.Document(docx_path)
        print(f"Successfully opened document with {len(doc.paragraphs)} paragraphs")
        
        # Print the first 20 paragraphs to see the structure
        for i, para in enumerate(doc.paragraphs[:50]):
            if para.text.strip():
                print(f"Paragraph {i}: {para.text}")
        
        return True
    except Exception as e:
        print(f"Error reading document: {e}")
        return False

if __name__ == "__main__":
    # Try both document files
    files_to_try = [
        "SEPM Report Part 2-prasheeth.docx",
        "NutriFit_TOC.docx"
    ]
    
    for file in files_to_try:
        if os.path.exists(file):
            print(f"\nAttempting to read: {file}")
            success = extract_toc_from_docx(file)
            if success:
                print(f"Successfully analyzed {file}")
            else:
                print(f"Failed to analyze {file}")
        else:
            print(f"File not found: {file}") 